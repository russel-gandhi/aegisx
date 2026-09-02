"""
Document parsing, chunking, and indexing pipeline (Phase 06.1, plan
06.1-01, RAG-01/RAG-02, D-01/D-04).

Ticket: n/a (roadmap phase 06.1) | Requirements: RAG-01, RAG-02
Source: AegisX-AI-Project-Bible-v6.md Section 9.1 (`chunk_text` sample);
Bible Section 15's hybrid-retrieval spec.

Mirrors `app.graph.evidence_graph`'s `build_graph`/`persist_graph`
two-phase shape: compute in pure Python first (`parse_text`,
`chunk_blocks` -- no I/O, no DB, no network), then one write phase
(`index_document`) that embeds and persists. Every SQL statement here uses
asyncpg `$N`-style placeholders exclusively -- `content`/`section` and
every other upload-derived value cross a trust boundary at
`routes/documents.py` (Security Domain V5, SQL injection).

Deviation 16: Bible Section 9.1's `chunk_text()` sample slices Python
string characters despite its own "Token-based chunking logic here"
comment (06.1-RESEARCH.md Pitfall 5) -- the sample's actual body is
`text[i:i + chunk_size]` over `range(0, len(text), chunk_size - overlap)`,
which is character-index chunking, not token-based chunking. This module
chunks on word boundaries deliberately, with no tokenizer dependency.

Plumbing note: `chunk_blocks()`'s frozen signature (`blocks -> List[Chunk]`)
does not receive the original upload filename, so each `Chunk.metadata`
carries only `{"chunker": "word-bounded", "chunk_words": CHUNK_WORDS}` at
this layer. `routes/documents.py` (plan 06.1-01 Task 2), which does have
the client filename, sets `metadata["source_filename"]` on each chunk
before calling `index_document()`.

Format-provenance note (plan 06.1-04, D-04): each of the four supported
formats yields a different, format-honest subset of `ParsedBlock`'s
`section`/`page` fields -- a `None` here is not a bug, it is what that
format actually supports:
- `pdf` -> real 1-indexed page numbers (`parse_pdf`) plus a heuristic
  heading-derived `section` (no reliable page-independent heading model
  exists for PDF, so this is a best-effort heuristic, not exact).
- `docx` -> real `Heading N`-style `section` text (`parse_docx`), never a
  `page` (DOCX carries no reliable page model at the byte level).
- `csv` -> `section` is the source filename stem and `page` is the
  1-indexed block number (`parse_csv`) -- a CSV has no prose headings or
  pages of its own, so these are the closest honest analogues.
- `text`/`markdown` -> real ATX (`#`..`######`) heading `section`
  (`parse_text`, plan 06.1-01), never a `page`.
"""

import csv
import io
import json
import logging
import os
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, FrozenSet, List, Optional, Tuple

from app.retrieval.embeddings import call_embeddings_batch
from app.retrieval.qdrant_store import upsert_chunks

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover -- pinned in requirements.txt
    docx = None  # type: ignore[assignment]

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover -- pinned in requirements.txt
    PdfReader = None  # type: ignore[assignment]

logger = logging.getLogger(__name__)

INGESTION_STAGES: Tuple[str, ...] = (
    "uploading",
    "parsing",
    "structure",
    "chunking",
    "indexing",
    "ready",
)

CHUNK_WORDS: int = 350
CHUNK_OVERLAP_WORDS: int = 50
MAX_UPLOAD_BYTES: int = 10 * 1024 * 1024
MAX_CHUNKS_PER_DOCUMENT: int = 500
SUPPORTED_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".markdown"})
MAX_PDF_PAGES: int = 500
MAX_CSV_ROWS: int = 5000
CSV_ROWS_PER_BLOCK: int = 25

# Frozen allowlist: format key -> the file extensions that claim it.
# `detect_format` resolves the extension through this map first, then
# confirms the payload's own bytes agree (T-06.1-21) -- the extension
# alone is never sufficient.
SUPPORTED_FORMATS: Dict[str, FrozenSet[str]] = {
    "text": frozenset({".txt", ".md", ".markdown"}),
    "pdf": frozenset({".pdf"}),
    "docx": frozenset({".docx"}),
    "csv": frozenset({".csv"}),
}

# Magic-byte signatures for the two binary formats. `text`/`csv` are
# confirmed by a strict UTF-8 decode instead (no fixed binary signature
# exists for either).
MAGIC_BYTES: Dict[str, bytes] = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",
}

_HEADING_RE_PREFIX = "#"


@dataclass
class ParsedBlock:
    text: str
    section: Optional[str]
    page: Optional[int]


@dataclass
class Chunk:
    chunk_id: str
    content: str
    section: Optional[str]
    page: Optional[int]
    chunk_index: int
    parent_chunk_id: Optional[str]
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class IngestResult:
    chunk_count: int
    indexed_vector_count: int
    status: str
    failed_stage: Optional[str] = None


def _match_heading(line: str) -> Optional[str]:
    """Return the heading text for a Markdown ATX heading line
    (`^#{1,6}\\s+(.*)$`), or `None` if `line` is not a heading."""
    stripped = line.lstrip()
    hashes = len(stripped) - len(stripped.lstrip(_HEADING_RE_PREFIX))
    if hashes < 1 or hashes > 6:
        return None
    rest = stripped[hashes:]
    if not rest[:1].isspace():
        return None
    return rest.strip()


def parse_text(raw: bytes, filename: str) -> List[ParsedBlock]:
    """Decode `raw` as UTF-8 (`errors="replace"` -- an upload must never
    crash the parser on an invalid byte sequence) and split on Markdown
    ATX headings (`#`..`######`) to assign each block's `section`. Plain
    `.txt` content with no headings falls through to a single block with
    `section=None`. `page` is always `None` for text/Markdown input --
    only the binary formats plan 06.1-04 adds carry real page numbers.
    """
    text = raw.decode("utf-8", errors="replace")
    lines = text.splitlines()

    blocks: List[ParsedBlock] = []
    current_section: Optional[str] = None
    current_lines: List[str] = []

    def _flush() -> None:
        content = "\n".join(current_lines).strip()
        if content:
            blocks.append(ParsedBlock(text=content, section=current_section, page=None))

    for line in lines:
        heading = _match_heading(line)
        if heading is not None:
            _flush()
            current_lines = []
            current_section = heading
        else:
            current_lines.append(line)
    _flush()

    return blocks


def _is_pdf_heading_line(stripped: str) -> bool:
    """Heading heuristic for `parse_pdf` (plan 06.1-04 Task 1 <action>):
    a candidate heading line is short (<=80 chars), non-empty, carries no
    terminal period, and does not begin with a lowercase letter -- the
    same test applied to every extracted line, never a guessed value."""
    return (
        bool(stripped)
        and len(stripped) <= 80
        and not stripped.endswith(".")
        and not stripped[0].islower()
    )


def _append_pdf_block(
    blocks: List[ParsedBlock], lines: List[str], section: Optional[str], page: int
) -> None:
    content = "\n".join(lines).strip()
    if content:
        blocks.append(ParsedBlock(text=content, section=section, page=page))


def parse_pdf(raw: bytes, filename: str) -> List[ParsedBlock]:
    """Extracts text from an in-memory PDF (`io.BytesIO(raw)` -- no temp
    file is ever written) using `pypdf`, one `ParsedBlock` per
    contiguous run of body lines. `page` is always the real 1-indexed
    page number the text came from. `section` is set by
    `_is_pdf_heading_line`'s heuristic; a page with no heading-shaped
    line yields blocks with `section=None` rather than a guessed value.

    Stops at `MAX_PDF_PAGES` (500) and logs the truncation (T-06.1-20).
    Never raises to its caller: an encrypted, corrupt, or otherwise
    unparseable PDF is logged and returns `[]`, which the calling route
    maps to `status="FAILED"`, `failed_stage="parsing"` (T-06.1-22).
    """
    if PdfReader is None:  # pragma: no cover -- pinned in requirements.txt
        logger.error("parse_pdf: pypdf is not installed")
        return []

    try:
        reader = PdfReader(io.BytesIO(raw))
        total_pages = len(reader.pages)
    except Exception:
        logger.warning("parse_pdf: failed to open %s as a PDF", filename, exc_info=True)
        return []

    if total_pages > MAX_PDF_PAGES:
        logger.warning(
            "parse_pdf: %s has %d pages, truncating to MAX_PDF_PAGES=%d",
            filename,
            total_pages,
            MAX_PDF_PAGES,
        )

    blocks: List[ParsedBlock] = []
    current_section: Optional[str] = None

    try:
        for page_num, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                logger.warning(
                    "parse_pdf: failed to extract text from page %d of %s",
                    page_num,
                    filename,
                    exc_info=True,
                )
                page_text = ""

            current_lines: List[str] = []
            for line in page_text.splitlines():
                stripped = line.strip()
                if _is_pdf_heading_line(stripped):
                    _append_pdf_block(blocks, current_lines, current_section, page_num)
                    current_lines = []
                    current_section = stripped
                else:
                    current_lines.append(line)
            _append_pdf_block(blocks, current_lines, current_section, page_num)
    except Exception:
        logger.warning("parse_pdf: unexpected error parsing %s", filename, exc_info=True)
        return []

    return blocks


_DOCX_BODY_TAG_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DOCX_PARAGRAPH_TAG = _DOCX_BODY_TAG_NS + "p"
_DOCX_TABLE_TAG = _DOCX_BODY_TAG_NS + "tbl"


def parse_docx(raw: bytes, filename: str) -> List[ParsedBlock]:
    """Extracts text from an in-memory DOCX (`io.BytesIO(raw)` -- no temp
    file is ever written) using `python-docx`. `section` is set from
    `paragraph.style.name.startswith("Heading")` paragraphs, using the
    heading's own text as the section title; `page` is always `None`
    (DOCX carries no reliable page model at the byte level). Table cell
    text is emitted as its own body block, joining each row's cells with
    `" | "` and rows with `"\\n"`.

    2026-09-01: walks `document.element.body`'s XML children directly
    (`<w:p>` and `<w:tbl>` elements, in true document order) instead of
    python-docx's separate `document.paragraphs`/`document.tables`
    collections. Those two collections are NOT interleaved -- iterating
    them one after the other (the previous approach) meant `current_section`
    had already advanced through every heading in the document by the time
    the tables loop ran, so every single table in the document silently
    got attached to whichever heading happened to be LAST in the file,
    regardless of which heading actually preceded that table. Confirmed
    live against the Dummy_data GxP corpus: a document's 3.1 role-matrix
    table and an unrelated training-curriculum table both came back tagged
    with the document's final section ("G3 Quality gate") instead of their
    real, structurally-correct sections. Walking the body in source order
    keeps `current_section` accurate for both paragraphs and tables alike.

    Never raises to its caller: a corrupt or non-OOXML payload is logged
    and returns `[]` (T-06.1-22), mirroring `parse_pdf`'s contract.
    """
    if docx is None:  # pragma: no cover -- pinned in requirements.txt
        logger.error("parse_docx: python-docx is not installed")
        return []

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception:
        logger.warning("parse_docx: failed to open %s as a DOCX", filename, exc_info=True)
        return []

    blocks: List[ParsedBlock] = []
    current_section: Optional[str] = None
    current_lines: List[str] = []

    def _flush() -> None:
        content = "\n".join(current_lines).strip()
        if content:
            blocks.append(ParsedBlock(text=content, section=current_section, page=None))

    try:
        # `Paragraph`/`Table` only need `parent` to resolve `.part` (for
        # relationship lookups like hyperlinks/images) -- `document` itself
        # satisfies that, so it's used uniformly for every child rather than
        # fishing a parent out of `document.paragraphs`/`document.tables`
        # (both fragile for a document that happens to have zero of either).
        for child in document.element.body.iterchildren():
            if child.tag == _DOCX_PARAGRAPH_TAG:
                paragraph = docx.text.paragraph.Paragraph(child, document)
                style_name = paragraph.style.name if paragraph.style is not None else ""
                if style_name.startswith("Heading"):
                    _flush()
                    current_lines = []
                    current_section = paragraph.text.strip()
                else:
                    text = paragraph.text.strip()
                    if text:
                        current_lines.append(text)
            elif child.tag == _DOCX_TABLE_TAG:
                _flush()
                current_lines = []
                table = docx.table.Table(child, document)
                row_lines = [
                    " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
                ]
                table_text = "\n".join(row_lines).strip()
                if table_text:
                    blocks.append(ParsedBlock(text=table_text, section=current_section, page=None))
        _flush()
    except Exception:
        logger.warning("parse_docx: unexpected error parsing %s", filename, exc_info=True)
        return []

    return blocks


def parse_csv(raw: bytes, filename: str) -> List[ParsedBlock]:
    """Structured row/column parsing (06.1-CONTEXT.md D-04) -- CSV takes
    its own dedicated path here, entirely independent of the ATX-heading
    prose chunker used for Markdown/plain-text input. Decodes UTF-8 with
    `errors="replace"`, reads via the stdlib `csv.DictReader`, and groups
    at most `MAX_CSV_ROWS` (5000) rows into blocks of `CSV_ROWS_PER_BLOCK`
    (25) data rows each. Each block's text repeats the header line at the
    top so a retrieved chunk is self-describing, and every data row is
    rendered as `"col: value"` pairs joined by `"; "` (not raw
    comma-separated text) so lexical search can match on a column name.
    `section` is the source filename stem; `page` is the 1-indexed block
    number -- CSV has no prose headings or pages of its own, so these are
    the closest honest analogues (see module docstring).

    A file with no data rows returns `[]` -- the caller (`routes/documents.py`)
    treats this as an honest `chunk_count=0`/`status="READY"` zero, never
    an error and never a fabricated row. Never raises to its caller: a
    malformed CSV payload is logged and returns `[]`, mirroring
    `parse_pdf`/`parse_docx`'s contract.
    """
    text = raw.decode("utf-8", errors="replace")

    try:
        reader = csv.DictReader(io.StringIO(text))
        fieldnames = reader.fieldnames or []
        rows: List[Dict[str, Any]] = []
        for row in reader:
            if len(rows) >= MAX_CSV_ROWS:
                logger.warning(
                    "parse_csv: %s exceeds MAX_CSV_ROWS=%d, truncating",
                    filename,
                    MAX_CSV_ROWS,
                )
                break
            rows.append(row)
    except Exception:
        logger.warning("parse_csv: failed to parse %s as CSV", filename, exc_info=True)
        return []

    if not rows or not fieldnames:
        return []

    section = os.path.splitext(os.path.basename(filename))[0]
    header_line = " | ".join(fieldnames)

    blocks: List[ParsedBlock] = []
    block_index = 0
    for start in range(0, len(rows), CSV_ROWS_PER_BLOCK):
        block_index += 1
        block_rows = rows[start : start + CSV_ROWS_PER_BLOCK]
        row_lines = [
            "; ".join(f"{col}: {value}" for col, value in row.items() if value)
            for row in block_rows
        ]
        block_text = "\n".join([header_line] + row_lines)
        blocks.append(ParsedBlock(text=block_text, section=section, page=block_index))

    return blocks


def _extension_format(filename: str) -> Optional[str]:
    extension = os.path.splitext(filename.lower())[1]
    for fmt, extensions in SUPPORTED_FORMATS.items():
        if extension in extensions:
            return fmt
    return None


def detect_format(raw: bytes, filename: str) -> Optional[str]:
    """Resolves `filename`'s extension to a format key via
    `SUPPORTED_FORMATS`, then confirms `raw`'s own bytes agree
    (T-06.1-21, file-type confusion): for `pdf`/`docx` the payload must
    start with the matching `MAGIC_BYTES` signature; for `csv`/`text` the
    first 4096 bytes must decode as strict UTF-8. Returns `None` on any
    mismatch or on an extension outside `SUPPORTED_FORMATS` -- the
    extension alone is never sufficient, and an unknown extension is
    rejected even when the underlying bytes happen to be valid text.
    """
    fmt = _extension_format(filename)
    if fmt is None:
        return None

    if fmt in MAGIC_BYTES:
        if not raw.startswith(MAGIC_BYTES[fmt]):
            return None
        return fmt

    try:
        raw[:4096].decode("utf-8")
    except UnicodeDecodeError:
        return None
    return fmt


def parse_document(raw: bytes, filename: str, fmt: str) -> List[ParsedBlock]:
    """Dispatches to the parser matching `fmt`. `fmt` is expected to
    already be a `detect_format`-validated key -- the route gates on
    `detect_format` before ever reaching this dispatcher, so an unknown
    key here indicates a caller bug, not untrusted input, hence the
    `ValueError` rather than a degrade-to-empty-list."""
    if fmt == "text":
        return parse_text(raw, filename)
    if fmt == "pdf":
        return parse_pdf(raw, filename)
    if fmt == "docx":
        return parse_docx(raw, filename)
    if fmt == "csv":
        return parse_csv(raw, filename)
    raise ValueError(f"parse_document: unknown format {fmt!r}")


def chunk_blocks(blocks: List[ParsedBlock]) -> List[Chunk]:
    """Word-bounded chunking (Deviation 16): `CHUNK_WORDS`-sized windows
    with `CHUNK_OVERLAP_WORDS` of overlap between consecutive chunks
    within the same block. `chunk_index` is a single document-global
    monotonic counter (not reset per block/section). `parent_chunk_id` is
    `None` for the first chunk produced from a given block (its section's
    leader) and that leader's `chunk_id` for every later chunk from the
    same block. Stops entirely at `MAX_CHUNKS_PER_DOCUMENT` -- a
    pathologically large upload is truncated, not rejected outright
    (rejection-by-size is `MAX_UPLOAD_BYTES`'s job, enforced by the route
    before parsing ever runs). Never returns an empty-content chunk.
    """
    chunks: List[Chunk] = []
    chunk_index = 0

    for block in blocks:
        words = block.text.split()
        if not words:
            continue

        section_leader_id: Optional[str] = None
        start = 0
        while start < len(words):
            if len(chunks) >= MAX_CHUNKS_PER_DOCUMENT:
                return chunks

            end = min(start + CHUNK_WORDS, len(words))
            content = " ".join(words[start:end])
            new_chunk_id = str(uuid.uuid4())
            parent_id = section_leader_id  # None for the section's first chunk

            chunks.append(
                Chunk(
                    chunk_id=new_chunk_id,
                    content=content,
                    section=block.section,
                    page=block.page,
                    chunk_index=chunk_index,
                    parent_chunk_id=parent_id,
                    metadata={"chunker": "word-bounded", "chunk_words": CHUNK_WORDS},
                )
            )
            if section_leader_id is None:
                section_leader_id = new_chunk_id
            chunk_index += 1

            if end == len(words):
                break
            start = end - CHUNK_OVERLAP_WORDS

    return chunks


async def index_document(
    pool: Any,
    qdrant_client: Any,
    document_id: str,
    system_id: str,
    chunks: List[Chunk],
) -> IngestResult:
    """Embed every chunk's content (one `call_embeddings_batch()` call,
    `task_type="RETRIEVAL_DOCUMENT"`), then write Postgres and Qdrant
    together inside a single Postgres transaction: if the Qdrant upsert
    raises, the `document_chunks` INSERTs already issued in this same
    transaction roll back with it, so a Qdrant failure never leaves an
    orphaned Postgres row with no matching vector.

    If ANY embedding response is degraded, writes nothing to either store
    and returns `IngestResult(status="FAILED", failed_stage="indexing")` --
    an honest failure envelope (D-09 discipline), never a zero-vector
    index and never a fabricated success.
    """
    if not chunks:
        return IngestResult(chunk_count=0, indexed_vector_count=0, status="READY", failed_stage=None)

    embeddings = await call_embeddings_batch(
        [chunk.content for chunk in chunks], task_type="RETRIEVAL_DOCUMENT"
    )

    if any(embedding.degraded for embedding in embeddings):
        logger.warning(
            "index_document: embedding provider degraded for document_id=%s "
            "(%d chunk(s)); writing nothing to Postgres or Qdrant.",
            document_id,
            len(chunks),
        )
        return IngestResult(chunk_count=0, indexed_vector_count=0, status="FAILED", failed_stage="indexing")

    async with pool.acquire() as conn:
        async with conn.transaction():
            for chunk, embedding in zip(chunks, embeddings):
                await conn.execute(
                    "INSERT INTO document_chunks "
                    "(chunk_id, document_id, content, embedding_id, section, page, "
                    "parent_chunk_id, chunk_index, metadata) "
                    "VALUES ($1::uuid, $2, $3, $4, $5, $6, $7::uuid, $8, $9::jsonb)",
                    chunk.chunk_id,
                    document_id,
                    chunk.content,
                    chunk.chunk_id,
                    chunk.section,
                    chunk.page,
                    chunk.parent_chunk_id,
                    chunk.chunk_index,
                    json.dumps(chunk.metadata),
                )

            points = [
                {
                    "chunk_id": chunk.chunk_id,
                    "document_id": document_id,
                    "system_id": system_id,
                    "vector": embedding.vector,
                    "section": chunk.section,
                    "page": chunk.page,
                    "chunk_index": chunk.chunk_index,
                }
                for chunk, embedding in zip(chunks, embeddings)
            ]
            indexed_count = await upsert_chunks(qdrant_client, points)

    return IngestResult(
        chunk_count=len(chunks),
        indexed_vector_count=indexed_count,
        status="READY",
        failed_stage=None,
    )
