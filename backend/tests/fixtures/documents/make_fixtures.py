"""
Deterministic fixture generator for the PDF/DOCX/CSV format-parsing
tests (Phase 06.1, plan 06.1-04, Task 1/Task 2, RAG-01/D-04).

Regenerates `sop_extract.pdf`, `validation_protocol.docx`, and
`traceability_matrix.csv` from synthetic GxP-flavoured text so the three
committed binaries are reproducible rather than opaque. Run directly:

    python tests/fixtures/documents/make_fixtures.py

`build_minimal_pdf()` hand-constructs a minimal valid PDF 1.4 byte
stream with no third-party PDF-authoring dependency: `reportlab` is not
present in `backend/requirements.txt`, and this repo's package-install
discipline (executor deviation Rule 3's exclusion for package-manager
installs) treats adding a new dependency as a checkpoint-worthy human
decision, not something a task takes unilaterally mid-execution. `pypdf`
itself (already pinned, plan 06.1-01) offers no simple high-level
text-drawing API on its `PdfWriter`, so this function writes the PDF
object graph (catalog, pages, a standard non-embedded Helvetica font,
and one content stream per page) directly, placing each input line with
its own `Tj` text-show operator so `pypdf.PageObject.extract_text()`
yields the exact same line-for-line structure this module's `pages`
argument describes -- `parse_pdf()`'s heading heuristic depends on that
line-for-line fidelity.

All content below is synthetic placeholder text describing no real
regulated system, procedure, or record -- nothing here is real
regulated content.
"""

import io
import os

PDF_PAGE_WIDTH = 612
PDF_PAGE_HEIGHT = 792


def _escape_pdf_text(text: str) -> str:
    """Escapes the three characters that are syntactically significant
    inside a PDF literal string (`(...)`)."""
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def build_minimal_pdf(pages):
    """Builds a minimal PDF 1.4 document from `pages`: a list of
    page-line-lists (each inner list is one page's lines, first line
    typically a heading candidate). Every line becomes its own `Tj`
    text-show operator at a fixed 14pt leading, so a reader that splits
    `extract_text()` on newlines recovers the same lines given here.
    """
    n_pages = len(pages)
    catalog_num = 1
    pages_num = 2
    font_num = 3
    page_obj_start = 4
    content_obj_start = page_obj_start + n_pages

    objects = []
    kids = " ".join(f"{page_obj_start + i} 0 R" for i in range(n_pages))
    objects.append((catalog_num, f"<< /Type /Catalog /Pages {pages_num} 0 R >>"))
    objects.append((pages_num, f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>"))
    objects.append((font_num, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

    for i, lines in enumerate(pages):
        page_num = page_obj_start + i
        content_num = content_obj_start + i
        objects.append(
            (
                page_num,
                f"<< /Type /Page /Parent {pages_num} 0 R "
                f"/MediaBox [0 0 {PDF_PAGE_WIDTH} {PDF_PAGE_HEIGHT}] "
                f"/Resources << /Font << /F1 {font_num} 0 R >> >> "
                f"/Contents {content_num} 0 R >>",
            )
        )
        stream_lines = ["BT", "/F1 12 Tf", "72 720 Td", "14 TL"]
        for j, line in enumerate(lines):
            text = _escape_pdf_text(line)
            if j > 0:
                stream_lines.append("T*")
            stream_lines.append(f"({text}) Tj")
        stream_lines.append("ET")
        stream_content = "\n".join(stream_lines)
        stream_bytes = stream_content.encode("latin-1")
        objects.append(
            (
                content_num,
                f"<< /Length {len(stream_bytes)} >>\nstream\n{stream_content}\nendstream",
            )
        )

    buf = io.BytesIO()
    buf.write(b"%PDF-1.4\n")
    offsets = {}
    max_obj = content_obj_start + n_pages - 1
    for num, body in objects:
        offsets[num] = buf.tell()
        buf.write(f"{num} 0 obj\n".encode("latin-1"))
        buf.write(body.encode("latin-1"))
        buf.write(b"\nendobj\n")

    xref_offset = buf.tell()
    buf.write(f"xref\n0 {max_obj + 1}\n".encode("latin-1"))
    buf.write(b"0000000000 65535 f \n")
    for num in range(1, max_obj + 1):
        buf.write(f"{offsets.get(num, 0):010d} 00000 n \n".encode("latin-1"))
    buf.write(f"trailer\n<< /Size {max_obj + 1} /Root {catalog_num} 0 R >>\n".encode("latin-1"))
    buf.write(f"startxref\n{xref_offset}\n%%EOF".encode("latin-1"))
    return buf.getvalue()


SOP_PAGES = [
    [
        "Standard Operating Procedure",
        "This document describes the standard procedure for equipment calibration.",
        "All steps below must be followed exactly as written in this synthetic fixture.",
    ],
    [
        "Equipment Calibration Requirements",
        "Calibration intervals are defined per instrument classification below.",
        "Records of every calibration event must be retained for audit purposes here.",
    ],
    [
        "Record Retention Policy",
        "All calibration records are retained for a minimum of seven years per policy.",
        "Retention periods align with applicable regulatory guidance documents cited.",
    ],
]


def build_sop_pdf() -> bytes:
    return build_minimal_pdf(SOP_PAGES)


def build_validation_protocol_docx() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Installation Qualification", level=1)
    document.add_paragraph(
        "This synthetic validation protocol documents the installation "
        "qualification steps for a demonstration GxP system."
    )
    document.add_paragraph(
        "Every installation step below must be executed and independently "
        "verified before proceeding to operational qualification."
    )

    document.add_heading("Operational Qualification", level=1)
    document.add_paragraph(
        "This section documents the operational qualification test cases "
        "exercised against the demonstration system."
    )

    table = document.add_table(rows=3, cols=3)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Test Step"
    header_cells[1].text = "Expected Result"
    header_cells[2].text = "Actual Result"
    table.rows[1].cells[0].text = "Power-on self-test"
    table.rows[1].cells[1].text = "System boots without error"
    table.rows[1].cells[2].text = "Pass"
    table.rows[2].cells[0].text = "Network connectivity check"
    table.rows[2].cells[1].text = "Ping succeeds within 1s"
    table.rows[2].cells[2].text = "Pass"

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_traceability_matrix_csv(row_count: int = 60) -> bytes:
    import csv as _csv

    buf = io.StringIO()
    writer = _csv.DictWriter(
        buf, fieldnames=["urs_id", "requirement", "test_case_id", "execution_status"]
    )
    writer.writeheader()
    for i in range(1, row_count + 1):
        writer.writerow(
            {
                "urs_id": f"URS-{i:03d}",
                "requirement": f"Synthetic requirement text for traceability row {i}.",
                "test_case_id": f"TC-{i:03d}",
                "execution_status": "PASS" if i % 5 else "FAIL",
            }
        )
    return buf.getvalue().encode("utf-8")


def main() -> None:
    fixtures_dir = os.path.dirname(os.path.abspath(__file__))

    sop_path = os.path.join(fixtures_dir, "sop_extract.pdf")
    with open(sop_path, "wb") as f:
        f.write(build_sop_pdf())
    print(f"wrote {sop_path}")

    docx_path = os.path.join(fixtures_dir, "validation_protocol.docx")
    with open(docx_path, "wb") as f:
        f.write(build_validation_protocol_docx())
    print(f"wrote {docx_path}")

    csv_path = os.path.join(fixtures_dir, "traceability_matrix.csv")
    with open(csv_path, "wb") as f:
        f.write(build_traceability_matrix_csv())
    print(f"wrote {csv_path}")


if __name__ == "__main__":
    main()
